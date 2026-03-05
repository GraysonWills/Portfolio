#!/usr/bin/env python3
"""
Generate Word companions for the comprehensive architecture documentation pack.

Outputs:
  - output/doc/portfolio-platform-comprehensive-documentation.docx
  - output/doc/portfolio-platform-visual-mockups.docx
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "comprehensive"
OUTPUT_DIR = DOCS_DIR / "word"

ALL_DOCS = [
    DOCS_DIR / "README.md",
    DOCS_DIR / "01-project-system-visual-mockup.md",
    DOCS_DIR / "02-aws-architecture-visual-mockup.md",
    DOCS_DIR / "03-backend-service-interplay.md",
    DOCS_DIR / "04-code-deep-dive.md",
    DOCS_DIR / "05-reference-matrix.md",
    DOCS_DIR / "06-authoring-hotkeys-and-page-checklist.md",
]

VISUAL_DOCS = [
    DOCS_DIR / "01-project-system-visual-mockup.md",
    DOCS_DIR / "02-aws-architecture-visual-mockup.md",
    DOCS_DIR / "03-backend-service-interplay.md",
]


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def _render_markdown(doc: Document, markdown_lines: Iterable[str]) -> None:
    in_code = False
    code_lang = ""

    for raw_line in markdown_lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip().lower()
                if code_lang:
                    _add_code_paragraph(doc, f"[code block: {code_lang}]")
            else:
                in_code = False
                code_lang = ""
                doc.add_paragraph("")
            continue

        if in_code:
            _add_code_paragraph(doc, line)
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            title = heading_match.group(2).strip()
            doc.add_heading(title, level=level)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(text, style="List Number")
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            _add_code_paragraph(doc, stripped)
            continue

        doc.add_paragraph(line)


def build_doc(output_path: Path, source_docs: List[Path], title: str) -> None:
    doc = Document()
    _set_default_font(doc)

    doc.add_heading(title, level=0)
    stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc.add_paragraph(f"Generated from Markdown sources on {stamp}.")
    doc.add_paragraph("")

    for idx, source in enumerate(source_docs):
        if not source.exists():
            continue
        doc.add_heading(source.name, level=1)
        doc.add_paragraph(str(source))
        doc.add_paragraph("")
        _render_markdown(doc, source.read_text(encoding="utf-8").splitlines())
        if idx < len(source_docs) - 1:
            doc.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    build_doc(
        output_path=OUTPUT_DIR / "portfolio-platform-comprehensive-documentation.docx",
        source_docs=ALL_DOCS,
        title="Portfolio Platform Comprehensive Documentation",
    )
    build_doc(
        output_path=OUTPUT_DIR / "portfolio-platform-visual-mockups.docx",
        source_docs=VISUAL_DOCS,
        title="Portfolio Platform Visual Mockups",
    )
    print("Generated DOCX files:")
    print(str(OUTPUT_DIR / "portfolio-platform-comprehensive-documentation.docx"))
    print(str(OUTPUT_DIR / "portfolio-platform-visual-mockups.docx"))


if __name__ == "__main__":
    main()
